from flask import Flask, jsonify, request,send_from_directory
from flask_cors import CORS
import psycopg2
import openai
from docx import Document
import os
import google.generativeai as genai
from dotenv import load_dotenv

app = Flask(__name__)
CORS(app)  # Allow frontend to access backend
# ✅ Define UPLOAD_FOLDER correctly

UPLOAD_FOLDER = "/tmp/recommendations"  # Use /tmp instead
os.makedirs(UPLOAD_FOLDER, exist_ok=True)  # Works on Vercel


# Connect to PostgreSQL
conn = psycopg2.connect(
    database="postgres",
    user="postgres",
    password="nishu*2003",
    host="localhost",
    port="5432"
)

UPLOAD_FOLDER = "/tmp/recommendations"  # Use /tmp instead of /var/task
os.makedirs(UPLOAD_FOLDER, exist_ok=True)  # Ensure the folder exists



cursor = conn.cursor()
load_dotenv()
GEMINI_API_KEY =  "AIzaSyBs-dLOxOuZ3WUNMIstaUnvxmAe4c20cFE"  # Direct assignment OR use .env file
genai.configure(api_key=GEMINI_API_KEY)

# Route to get all products
@app.route('/products', methods=['GET'])
def get_products():
    cursor.execute("SELECT product_id, name, price, category FROM Product;")
    products = cursor.fetchall()
    
    # Convert raw tuples into dictionaries
    product_list = [
        {"product_id": p[0], "name": p[1], "price": p[2], "category": p[3]} 
        for p in products
    ]
    
    return jsonify(product_list)
@app.route('/customers', methods=['GET'])
def get_customers():
    cursor.execute("SELECT customer_id, name, email, location FROM Customer;")
    customers = cursor.fetchall()
    
    customer_list = [
        {"customer_id": c[0], "name": c[1], "email": c[2], "location": c[3]} 
        for c in customers
    ]
    
    return jsonify(customer_list)
@app.route('/add-customer', methods=['POST'])
def add_customer():
    data = request.json
    name = data.get('name')
    email = data.get('email')
    location = data.get('location')

    if not name or not email or not location:
        return jsonify({"error": "All fields are required"}), 400

    try:
        cursor.execute("INSERT INTO Customer (name, email, location) VALUES (%s, %s, %s) RETURNING customer_id;",
                       (name, email, location))
        conn.commit()
        customer_id = cursor.fetchone()[0]
        return jsonify({"message": "Customer added successfully!", "customer_id": customer_id}), 201
    except Exception as e:
        conn.rollback()
        return jsonify({"error": str(e)}), 500
# Route to get all orders
@app.route('/orders', methods=['GET'])
def get_orders():
    cursor.execute("""
        SELECT o.order_id, c.name AS customer_name, p.name AS product_name, o.quantity, o.order_date
        FROM OrderTable o
        JOIN Customer c ON o.customer_id = c.customer_id
        JOIN Product p ON o.product_id = p.product_id;
    """)
    orders = cursor.fetchall()
    
    order_list = [
        {
            "order_id": o[0],
            "customer_name": o[1],
            "product_name": o[2],
            "quantity": o[3],
            "order_date": o[4]
        }
        for o in orders
    ]
    
    return jsonify(order_list)
@app.route('/place-order', methods=['POST'])
def place_order():
    data = request.json
    customer_id = data.get('customer_id')
    product_id = data.get('product_id')
    quantity = data.get('quantity')

    if not customer_id or not product_id or not quantity:
        return jsonify({"error": "All fields are required"}), 400

    try:
        cursor.execute(
            "INSERT INTO OrderTable (customer_id, product_id, quantity, order_date) VALUES (%s, %s, %s, NOW()) RETURNING order_id;",
            (customer_id, product_id, quantity)
        )
        conn.commit()
        order_id = cursor.fetchone()[0]
        return jsonify({"message": "Order placed successfully!", "order_id": order_id}), 201
    except Exception as e:
        conn.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/add-product', methods=['POST'])
def add_product():
    data = request.json
    name = data.get('name')
    price = data.get('price')
    category = data.get('category')

    if not name or not price or not category:
        return jsonify({"error": "All fields are required"}), 400

    try:
        cursor.execute("INSERT INTO Product (name, price, category) VALUES (%s, %s, %s) RETURNING product_id;",
                       (name, price, category))
        conn.commit()
        product_id = cursor.fetchone()[0]
        return jsonify({"message": "Product added successfully!", "product_id": product_id}), 201
    except Exception as e:
        conn.rollback()
        return jsonify({"error": str(e)}), 500

# Route to get a customer's purchase history
@app.route('/purchase-history/<int:customer_id>', methods=['GET'])
def get_purchase_history(customer_id):
    cursor.execute("""
        SELECT p.name, p.price, o.quantity, o.order_date
        FROM OrderTable o
        JOIN Product p ON o.product_id = p.product_id
        WHERE o.customer_id = %s;
    """, (customer_id,))
    history = cursor.fetchall()
    return jsonify(history)
# Function to fetch customer's purchase history
def get_customer_history(customer_id):
    cursor.execute("""
    SELECT p.name  
    FROM ordertable o
    JOIN Product p ON o.product_id = p.product_id
    WHERE o.customer_id = %s;
""", (customer_id,))
    return [row[0] for row in cursor.fetchall()]
def get_product_recommendations(history):
    prompt = f"""
    The customer has previously purchased: {', '.join(history)}.
    Based on this history, recommend 3 additional products they might like, along with a short explanation for each.
    """
    model = genai.GenerativeModel("gemini-1.5-pro")
    response = model.generate_content(prompt)

    return response.text if response.text else "No recommendations available."


# Function to write recommendations to a Word file
def write_to_word(customer_id, history, recommendations):
    doc = Document()
    doc.add_heading(f"Product Recommendations for Customer {customer_id}", level=1)

    doc.add_heading("Purchase History", level=2)
    for item in history:
        doc.add_paragraph(f"- {item}")

    doc.add_heading("Recommended Products", level=2)
    doc.add_paragraph(recommendations)

    filename = f"recommendations_customer_{customer_id}.docx"
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    doc.save(file_path)
    return filename
def get_db_cursor():
    return conn.cursor()

@app.route('/recommendations/<int:customer_id>')
def recommendations(customer_id):
    history = get_customer_history(customer_id)
    if not history:
        return jsonify({"error": "No purchase history found"}), 404

    recommendations = get_product_recommendations(history)
    word_file = write_to_word(customer_id, history, recommendations)

    return jsonify({
        "customer_id": customer_id,
        "purchase_history": history,
        "recommendations": recommendations,
        "word_file": word_file
    })

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
